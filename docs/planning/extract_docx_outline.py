from __future__ import annotations

import json
import re
import sys
from pathlib import Path

from docx import Document


def _normalize_ws(value: str) -> str:
    return re.sub(r"\s+", " ", (value or "").strip())


def _collect_headings(doc: Document) -> list[dict]:
    items: list[dict] = []
    for p in doc.paragraphs:
        text = _normalize_ws(p.text)
        if not text:
            continue
        style = (p.style.name or "") if p.style else ""
        style_lower = style.lower()
        if style_lower.startswith("heading"):
            m = re.search(r"(\d+)", style)
            level = int(m.group(1)) if m else 1
            items.append({"level": level, "text": text, "style": style})
    if items:
        return items

    num_re = re.compile(r"^(\d+(?:\.\d+)*)(?:\s+|\u3000+)(.+)$")
    for p in doc.paragraphs:
        text = _normalize_ws(p.text)
        if not text:
            continue
        m = num_re.match(text)
        if not m:
            continue
        num = m.group(1)
        items.append({"level": num.count(".") + 1, "text": text, "style": "numeric"})
    return items


def _collect_samples(doc: Document, limit: int) -> list[dict]:
    samples: list[dict] = []
    for p in doc.paragraphs:
        text = _normalize_ws(p.text)
        if not text:
            continue
        style = (p.style.name or "") if p.style else ""
        samples.append({"style": style, "text": text})
        if len(samples) >= limit:
            break
    return samples


def main() -> int:
    if len(sys.argv) < 2:
        print("Usage: python extract_docx_outline.py <docx_path> [sample_limit]", file=sys.stderr)
        return 2

    docx_path = Path(sys.argv[1]).expanduser().resolve()
    sample_limit = int(sys.argv[2]) if len(sys.argv) > 2 else 60
    doc = Document(str(docx_path))
    props = doc.core_properties

    result = {
        "file": {
            "name": docx_path.name,
            "path": str(docx_path),
        },
        "properties": {
            "title": props.title,
            "author": props.author,
            "last_modified_by": props.last_modified_by,
            "created": str(props.created) if props.created else None,
            "modified": str(props.modified) if props.modified else None,
        },
        "stats": {
            "paragraphs_total": len(doc.paragraphs),
            "paragraphs_nonempty": sum(1 for p in doc.paragraphs if _normalize_ws(p.text)),
            "tables_total": len(doc.tables),
        },
        "headings": _collect_headings(doc),
        "samples": _collect_samples(doc, sample_limit),
    }

    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

